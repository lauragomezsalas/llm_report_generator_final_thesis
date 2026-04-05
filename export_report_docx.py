import json
import os
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PRIMARY_COLOR = RGBColor(31, 78, 121)   # dark blue
TEXT_COLOR = RGBColor(34, 34, 34)
MUTED_COLOR = RGBColor(102, 102, 102)


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def set_default_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_COLOR

    document.styles["Title"].font.name = "Aptos Display"
    document.styles["Title"].font.size = Pt(24)
    document.styles["Title"].font.bold = True
    document.styles["Title"].font.color.rgb = PRIMARY_COLOR

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = PRIMARY_COLOR


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("Consulting AI Report | Page ")
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED_COLOR

    add_page_number(p)


def add_title_page(document: Document, report_title: str = "Strategic Consulting Report") -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(120)

    run = p.add_run(report_title)
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Prepared by the AI Consulting Workflow")
    run2.font.name = "Aptos"
    run2.font.size = Pt(11)
    run2.font.color.rgb = MUTED_COLOR

    document.add_page_break()


def add_section_heading(document: Document, text: str, level: int = 1):
    p = document.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)
    return p


def add_body_paragraph(document: Document, text: str, first: bool = False):
    if not text:
        return

    p = document.add_paragraph(style="Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if first:
        p.paragraph_format.space_before = Pt(2)
    p.add_run(str(text))
    return p


def add_label_value_paragraph(document: Document, label: str, value: str):
    if not value:
        return
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(4)

    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.name = "Aptos"
    r1.font.size = Pt(11)

    r2 = p.add_run(str(value))
    r2.font.name = "Aptos"
    r2.font.size = Pt(11)


def add_bullet(document: Document, text: str):
    if not text:
        return
    p = document.add_paragraph(style="Normal")
    p.style = document.styles["Normal"]
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("• ").bold = True
    p.add_run(str(text))


def add_divider(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—" * 28)
    run.font.color.rgb = MUTED_COLOR
    run.font.size = Pt(10)


def format_evidence_ids(evidence_ids):
    if not evidence_ids:
        return None
    return ", ".join(evidence_ids)


def add_numbered_heading(document, number: str, text: str, level: int = 1):
    p = document.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{number}. {text}" if level == 1 else f"{number} {text}")
    run.bold = True
    return p


def add_body_paragraph(document, text: str):
    if not text:
        return
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    p.add_run(str(text))


def add_bullet_list(document, items):
    if not items:
        return
    for item in items:
        p = document.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(3)
        p.add_run("• ").bold = True
        p.add_run(str(item))


def add_label_value_paragraph(document, label: str, value: str):
    if not value:
        return
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r2 = p.add_run(str(value))


def add_report_section(document, report_data: dict):
    if not isinstance(report_data, dict) or not report_data:
        add_numbered_heading(document, "1", "Report", level=1)
        add_body_paragraph(document, "No report data available.")
        return

    add_numbered_heading(document, "1", "Executive Summary", level=1)
    add_body_paragraph(document, report_data.get("executive_summary", ""))

    add_numbered_heading(document, "2", "Key Insights", level=1)
    add_bullet_list(document, report_data.get("key_insights", []))

    add_numbered_heading(document, "3", "Company and Market Overview", level=1)
    add_body_paragraph(document, report_data.get("company_and_market_overview", ""))

    add_numbered_heading(document, "4", "Strategic Alternatives", level=1)
    alternatives = report_data.get("strategic_alternatives_section", [])
    for idx, alt in enumerate(alternatives, start=1):
        add_numbered_heading(document, f"4.{idx}", alt.get("title", f"Alternative {idx}"), level=2)
        add_label_value_paragraph(document, "Alternative ID", alt.get("id", ""))
        add_label_value_paragraph(document, "Strategic rationale", alt.get("strategic_rationale", ""))
        add_label_value_paragraph(document, "Expected impact", alt.get("expected_impact_summary", ""))
        add_label_value_paragraph(document, "Risks", alt.get("risk_summary", ""))

        citations = alt.get("apa_citations", [])
        if citations:
            add_label_value_paragraph(document, "Sources", "; ".join(citations))

    add_numbered_heading(document, "5", "Trade-off Discussion", level=1)
    add_body_paragraph(document, report_data.get("trade_off_discussion", ""))

    add_numbered_heading(document, "6", "Financial Impact Summary", level=1)
    for item in report_data.get("financial_impact_summary", []):
        add_label_value_paragraph(document, item.get("metric", "Metric"), item.get("estimate", ""))
        add_body_paragraph(document, item.get("rationale", ""))
        citations = item.get("apa_citations", [])
        if citations:
            add_label_value_paragraph(document, "Sources", "; ".join(citations))

    add_numbered_heading(document, "7", "Final Recommendation", level=1)
    final_rec = report_data.get("final_recommendation", {})
    add_label_value_paragraph(document, "Selected alternative", final_rec.get("selected_alternative", ""))
    add_label_value_paragraph(document, "Justification", final_rec.get("justification", ""))
    add_label_value_paragraph(document, "Roadmap summary", final_rec.get("implementation_roadmap_summary", ""))
    citations = final_rec.get("apa_citations", [])
    if citations:
        add_label_value_paragraph(document, "Sources", "; ".join(citations))

    add_numbered_heading(document, "8", "Implementation Timeline", level=1)
    for idx, phase in enumerate(report_data.get("implementation_timeline", []), start=1):
        add_numbered_heading(
            document,
            f"8.{idx}",
            f'{phase.get("phase_title", f"Phase {idx}")} ({phase.get("timeline", "")})',
            level=2,
        )
        add_label_value_paragraph(document, "Objectives", "")
        add_bullet_list(document, phase.get("objectives", []))
        add_label_value_paragraph(document, "Key actions", "")
        add_bullet_list(document, phase.get("key_actions", []))
        add_label_value_paragraph(document, "Expected outputs", "")
        add_bullet_list(document, phase.get("expected_outputs", []))

    add_numbered_heading(document, "9", "Risks and Mitigation", level=1)
    for idx, item in enumerate(report_data.get("risks_and_mitigation", []), start=1):
        add_numbered_heading(document, f"9.{idx}", item.get("risk", f"Risk {idx}"), level=2)
        add_label_value_paragraph(document, "Mitigation", item.get("mitigation", ""))
        add_label_value_paragraph(document, "Severity", item.get("severity", ""))
        citations = item.get("apa_citations", [])
        if citations:
            add_label_value_paragraph(document, "Sources", "; ".join(citations))

    add_numbered_heading(document, "10", "Conclusion", level=1)
    add_body_paragraph(document, report_data.get("conclusion", ""))

    add_numbered_heading(document, "11", "References", level=1)
    refs = report_data.get("references", [])
    for ref in refs:
        p = document.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(str(ref))


def build_docx_from_run_record(run_record, output_path):
    document = Document()
    set_page_layout(document)
    set_default_styles(document)
    add_footer(document)

    report = run_record.get("report") or {}
    report_title = "Strategic Consulting Report"

    add_title_page(document, report_title="Strategic Consulting Report")    
    add_report_section(document, report)

    document.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 3:
        print("Usage: python export_report_docx.py <run_record_json_path> <output_docx_path>")
        sys.exit(1)

    input_json_path = sys.argv[1]
    output_docx_path = sys.argv[2]

    if not os.path.exists(input_json_path):
        print(f"Error: input file not found -> {input_json_path}")
        sys.exit(1)

    with open(input_json_path, "r", encoding="utf-8") as f:
        run_record = json.load(f)

    build_docx_from_run_record(run_record, output_docx_path)
    print(f"Report exported successfully to: {output_docx_path}")


if __name__ == "__main__":
    main()